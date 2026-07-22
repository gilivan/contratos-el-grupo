#!/usr/bin/env python3
"""
generar_contrato.py — Generador de documentos v2
Tipos: proveedor, freefan, acta_cierre
"""

import json
import sys
import copy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX

ASSETS_DIR = Path(__file__).parent / "assets"

# ── Datos fijos del CONTRATISTA ──────────────────────────────────────────────
# IMPORTANTE: actualizar CC_CONTRATISTA con la cédula real de Santiago Vélez Gracián
CONTRATISTA = {
    "nombre":   "SANTIAGO VÉLEZ GRACIÁN",
    "cc":       "CC_CONTRATISTA",   # <-- actualizar con la cédula real
    "telefono": "3154719586",
    "correo":   "info@elgrupo.com.co",
}

# ── Datos de las empresas contratistas ───────────────────────────────────────
COMPANY_DATA = {
    "fiera": {
        "razon_social": "FIERA S.A.S.",
        "nit":          "900.072.392-5",
    },
    "perez_villa": {
        "razon_social": "PÉREZ Y VILLA S.A.S.",
        "nit":          "890.926.395-8",
    },
}

# ── Plantillas ────────────────────────────────────────────────────────────────
TEMPLATE_FILES = {
    "fiera":       "plantilla_fiera.docx",
    "perez_villa": "plantilla_perez_villa.docx",
    "freefan":     "plantilla_freefan.docx",
    "acta_cierre": "plantilla_acta_cierre.docx",
}


# ─────────────────────────────────────────────────────────────────────────────
#  UTILIDADES
# ─────────────────────────────────────────────────────────────────────────────

def replace_and_highlight(paragraph, replacements):
    """
    Reemplaza placeholders en un párrafo y resalta en AMARILLO el texto sustituido.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if not any(key in full_text for key in replacements):
        return

    # Construir segmentos (texto, es_reemplazo)
    segments = [(full_text, False)]
    for key, value in replacements.items():
        new_segs = []
        for seg_text, seg_hi in segments:
            if not seg_hi and key in seg_text:
                parts = seg_text.split(key)
                for i, part in enumerate(parts):
                    if part:
                        new_segs.append((part, False))
                    if i < len(parts) - 1:
                        new_segs.append((str(value), True))
            else:
                new_segs.append((seg_text, seg_hi))
        segments = new_segs

    if not paragraph.runs or not segments:
        return

    # Capturar formato base del primer run
    base = paragraph.runs[0]
    base_bold      = base.bold
    base_italic    = base.italic
    base_underline = base.underline
    base_font_size = base.font.size
    base_font_name = base.font.name
    try:
        base_color = base.font.color.rgb if base.font.color.type else None
    except Exception:
        base_color = None

    def apply_format(run, highlighted):
        if base_bold is not None:
            run.bold = base_bold
        if base_italic is not None:
            run.italic = base_italic
        if base_underline is not None:
            run.underline = base_underline
        if base_font_size:
            run.font.size = base_font_size
        if base_font_name:
            run.font.name = base_font_name
        if base_color:
            try:
                run.font.color.rgb = base_color
            except Exception:
                pass
        if highlighted:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Limpiar runs existentes
    base.text = ""
    for run in paragraph.runs[1:]:
        run.text = ""

    # Reconstruir
    first = True
    for seg_text, seg_hi in segments:
        if first:
            base.text = seg_text
            if seg_hi:
                base.font.highlight_color = WD_COLOR_INDEX.YELLOW
            first = False
        else:
            new_run = paragraph.add_run(seg_text)
            apply_format(new_run, seg_hi)


def replace_in_doc(doc, replacements):
    """Aplica replacements en todos los párrafos del cuerpo y en todas las celdas de tablas."""
    for para in doc.paragraphs:
        replace_and_highlight(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_and_highlight(para, replacements)


def insert_activity_paragraphs(paragraph, activities):
    """
    Reemplaza el párrafo placeholder por un párrafo por actividad (• ítem, resaltado amarillo).
    """
    parent  = paragraph._p.getparent()
    para_idx = list(parent).index(paragraph._p)
    original_p = paragraph._p

    for activity in reversed(activities):
        new_p = copy.deepcopy(original_p)
        for r in new_p.findall(qn("w:r")):
            new_p.remove(r)

        new_r   = OxmlElement("w:r")
        orig_runs = original_p.findall(qn("w:r"))
        if orig_runs:
            orig_rpr = orig_runs[0].find(qn("w:rPr"))
            new_rpr  = copy.deepcopy(orig_rpr) if orig_rpr is not None else OxmlElement("w:rPr")
        else:
            new_rpr = OxmlElement("w:rPr")

        highlight_el = OxmlElement("w:highlight")
        highlight_el.set(qn("w:val"), "yellow")
        new_rpr.append(highlight_el)
        new_r.append(new_rpr)

        new_t = OxmlElement("w:t")
        new_t.text = f"• {activity}"
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        new_r.append(new_t)

        new_p.append(new_r)
        parent.insert(para_idx, new_p)

    parent.remove(original_p)


# ─────────────────────────────────────────────────────────────────────────────
#  CONTRATO PROVEEDOR  (Fiera / Pérez y Villa)
# ─────────────────────────────────────────────────────────────────────────────

def process_body(doc, data):
    """Reemplazos en el cuerpo del contrato proveedor."""
    body_replacements = {
        "[Cantidad de años o meses]":           data["duracion_cantidad"],
        "[Mes o Año]":                          data["duracion_unidad"],
        "[Fecha]":                              data["fecha_inicio"],
        "[Diligenciar según propuesta]":        data["valor_pago"],
        "[Concepto, valor y la periodicidad del pago]": data["valor_pago"],
        "[numero]":                             data["firma_dia"],
        "[nombre del mes]":                     data["firma_mes"],
        "[últimos dígitos del año]":            data["firma_anio"],
        "[NOMBRE CONTRANTE]":                   data["nombre_rl"],
        "[CÉDULA RL]":                          data["cedula_rl"],
        "[EMPRESA CONTRATANTE]":                data["empresa_contratante"],
    }

    for para in doc.paragraphs:
        if "[Actividades de la propuesta, se deben poner por ítem]" in para.text:
            insert_activity_paragraphs(para, data["actividades"])
        else:
            replace_and_highlight(para, body_replacements)


def process_table(doc, data, empresa):
    """Reemplazos en la tabla resumen de partes (contrato proveedor)."""
    if not doc.tables:
        return
    table = doc.tables[0]

    # Fila 0 — Empresa CONTRATANTE
    cell_empresa = table.rows[0].cells[1]
    if empresa == "fiera":
        replace_and_highlight(cell_empresa.paragraphs[0],
            {"[NOMBRE CONTRATANTE]": data["empresa_contratante"]})
    else:
        replace_and_highlight(cell_empresa.paragraphs[0],
            {"[NOMBRE EMPRESA CONTRATANTE]": data["empresa_contratante"]})

    # Fila 1 — Representante Legal + Cédula
    cell_rl = table.rows[1].cells[1]
    for para in cell_rl.paragraphs:
        replace_and_highlight(para, {
            "[NOMBRE REPRESENTANTE LEGAL CONTRATANTE]": data["nombre_rl"],
            "[CÉDULA RL]":                              data["cedula_rl"],
            "[NÚMERO CÉDULA RL]":                       data["cedula_rl"],
        })

    # Fila 2 — Dirección
    cell_dir = table.rows[2].cells[1]
    for para in cell_dir.paragraphs:
        replace_and_highlight(para, {
            "[DIRECCIÓN]":                data["direccion"],
            "[NOMENCLATURA DIRECCIÓN]":   data["direccion"],
        })

    # Fila 3 — NIT
    cell_nit = table.rows[3].cells[1]
    for para in cell_nit.paragraphs:
        replace_and_highlight(para, {
            "[Número del NIT del contratante]": data["nit"],
        })

    # Fila 4 — Contactos
    cell_contacto = table.rows[4].cells[1]
    for para in cell_contacto.paragraphs:
        replace_and_highlight(para, {
            "[NOMBRE CONTRATISTA]":    CONTRATISTA["nombre"],
            "[TELEFONO CONTRATISTA]":  CONTRATISTA["telefono"],
            "[CORREO CONTRATISTA]":    CONTRATISTA["correo"],
            "[NOMBRE CONTRATANTE]":    data["contacto_contratante_nombre"],
            "[TELEFONO CONTRATANTE]":  data["contacto_contratante_telefono"],
            "[CORREO CONTRATANTE]":    data["contacto_contratante_correo"],
        })


# ─────────────────────────────────────────────────────────────────────────────
#  CONTRATO FREEFAN
# ─────────────────────────────────────────────────────────────────────────────

def process_freefan(doc, data):
    """Reemplaza variables en la plantilla Freefan."""
    empresa  = data["empresa"]
    company  = COMPANY_DATA[empresa]

    replacements = {
        "[NUMERO_CONTRATO_FREEFAN]":  data["numero_contrato"],
        "[RAZON_SOCIAL]":             company["razon_social"],
        "[NIT_RAZON_SOCIAL]":         company["nit"],
        "[CONTRATISTA_FREEFAN]":      data["contratista_nombre"],
        "[CC_CONTRATISTA_FREEFAN]":   data["contratista_cc"],
        "[FECHA_INICIO_CONTRATO]":    data["fecha_inicio"],
        "[FECHA_FIN_CONTRATO]":       data["fecha_fin"],
        "[DURACIÓN_DÍAS_CONTRATO]":   data["duracion_dias"],
        "[VALOR_DEL_CONTRATO]":       data["valor"],
        "[DIA_FIRMA]":                data["firma_dia"],
        "[MES_FIRMA]":                data["firma_mes"],
        "[AÑO_FIRMA]":                data["firma_anio"],
    }

    replace_in_doc(doc, replacements)


# ─────────────────────────────────────────────────────────────────────────────
#  ACTA DE CIERRE Y TERMINACIÓN
# ─────────────────────────────────────────────────────────────────────────────

def process_acta(doc, data):
    """Reemplaza variables en el Acta de Cierre y Terminación."""
    empresa = data["empresa"]
    company = COMPANY_DATA[empresa]

    replacements = {
        "[NUMERO_DE_ACTA]":                 data["numero_acta"],
        "[FECHA_ACTUAL]":                   data["fecha_actual"],
        "[RAZON_SOCIAL]":                   company["razon_social"],
        "[NIT_RAZON_SOCIAL]":               company["nit"],
        "[CONTRATISTA_FREEFAN]":            data["contratista_nombre"],
        "[CC_CONTRATISTA_FREEFAN]":         data["contratista_cc"],
        "[FECHA_CONTRATO_FREEFAN]":         data["fecha_contrato"],
        "[FECHA_DE_FINALIZACIÓN_CONTRATO]": data["fecha_finalizacion"],
        "[NUMERO_CONTRATO_FREEFAN]":        data["numero_contrato"],
        "[VALOR_DE_CONTRATO]":              data["valor"],
    }

    replace_in_doc(doc, replacements)


# ─────────────────────────────────────────────────────────────────────────────
#  MAIN (uso por CLI)
# ─────────────────────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) != 3:
        print("Uso: python3 generar_contrato.py <datos.json> <output_path>")
        sys.exit(1)

    with open(sys.argv[1], "r", encoding="utf-8") as f:
        data = json.load(f)

    tipo    = data.get("tipo", "proveedor").lower().strip()
    empresa = data.get("empresa", "").lower().strip()

    if tipo == "proveedor":
        if empresa not in ("fiera", "perez_villa"):
            print(f"Error: empresa debe ser 'fiera' o 'perez_villa'")
            sys.exit(1)
        template_path = ASSETS_DIR / TEMPLATE_FILES[empresa]
        doc = Document(str(template_path))
        process_body(doc, data)
        process_table(doc, data, empresa)

    elif tipo == "freefan":
        if empresa not in COMPANY_DATA:
            print(f"Error: empresa debe ser 'fiera' o 'perez_villa'")
            sys.exit(1)
        template_path = ASSETS_DIR / TEMPLATE_FILES["freefan"]
        doc = Document(str(template_path))
        process_freefan(doc, data)

    elif tipo == "acta_cierre":
        if empresa not in COMPANY_DATA:
            print(f"Error: empresa debe ser 'fiera' o 'perez_villa'")
            sys.exit(1)
        template_path = ASSETS_DIR / TEMPLATE_FILES["acta_cierre"]
        doc = Document(str(template_path))
        process_acta(doc, data)

    else:
        print(f"Error: tipo '{tipo}' no reconocido. Use: proveedor, freefan, acta_cierre")
        sys.exit(1)

    if not template_path.exists():
        print(f"Error: plantilla no encontrada en {template_path}")
        sys.exit(1)

    doc.save(sys.argv[2])
    print(f"✓ Documento generado: {sys.argv[2]}")


if __name__ == "__main__":
    main()
